from docx import Document
from typing import Dict, List, Tuple
import re
from datetime import datetime


class OlympiadDocParser:
    """Парсер .docx файлов с кодами олимпиады"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = Document(file_path)
        
    def parse(self) -> Dict:
        """
        Основной метод парсинга документа
        
        Returns:
            Dict с ключами:
            - subject: название предмета
            - grade8_codes: список словарей {full_name, code}
            - grade9_codes: список кодов для 9 класса
        """
        print(f"\n{'='*60}")
        print(f"DEBUG: Начинаем парсинг файла: {self.file_path}")
        print(f"{'='*60}\n")
        
        result = {
            "subject": None,
            "grade8_codes": [],
            "grade9_codes": []
        }
        
        # Парсим предмет из текста
        result["subject"] = self._extract_subject()
        print(f"DEBUG: Определен предмет: {result['subject']}\n")
        
        # Парсим таблицы
        tables = self.document.tables
        print(f"DEBUG: Найдено таблиц в документе: {len(tables)}\n")
        
        if len(tables) > 0:
            print("DEBUG: Парсинг первой таблицы (8 класс)...")
            # Первая таблица - 8 класс с именами
            result["grade8_codes"] = self._parse_grade8_table(tables[0])
        
        # Парсим коды 9 класса из текста после таблицы
        print("\nDEBUG: Поиск кодов 9 класса...")
        result["grade9_codes"] = self._extract_grade9_codes()
        
        print(f"\n{'='*60}")
        print(f"DEBUG: Парсинг завершен:")
        print(f"  - Предмет: {result['subject']}")
        print(f"  - Кодов 8 класса: {len(result['grade8_codes'])}")
        print(f"  - Кодов 9 класса: {len(result['grade9_codes'])}")
        print(f"{'='*60}\n")
        
        return result
    
    def _extract_subject(self) -> str:
        """Извлекает название предмета из документа"""
        # Ищем в параграфах упоминание предмета
        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()
            # Паттерн для поиска предмета
            if text and any(word in text.lower() for word in ['физика', 'математика', 'химия', 'биология', 'информатика']):
                # Извлекаем название предмета
                match = re.search(r'(Физика|Математика|Химия|Биология|Информатика)', text, re.IGNORECASE)
                if match:
                    return match.group(1).capitalize()
        
        # Если не нашли в параграфах, проверяем заголовки таблиц
        for table in self.document.tables:
            if len(table.rows) > 0:
                first_row = table.rows[0]
                for cell in first_row.cells:
                    text = cell.text.strip()
                    match = re.search(r'(Физика|Математика|Химия|Биология|Информатика)', text, re.IGNORECASE)
                    if match:
                        return match.group(1).capitalize()
        
        return "Неизвестный предмет"
    
    def _parse_grade8_table(self, table) -> List[Dict]:
        """
        Парсит таблицу с кодами для 8 класса
        
        Ожидаемая структура таблицы:
        № | ФИО | Код
        """
        grade8_codes = []
        
        # Ищем строку с заголовками чтобы определить начало данных
        header_row_index = 0
        for i, row in enumerate(table.rows):
            cells = row.cells
            # Ищем строку где есть "№" или "ФИО" или название предмета
            row_text = ''.join(cell.text for cell in cells).lower()
            if '№' in row_text or 'фио' in row_text or 'физика' in row_text or 'математика' in row_text:
                header_row_index = i
                break
        
        print(f"DEBUG: Найден заголовок таблицы в строке {header_row_index}")
        
        # Начинаем со строки после заголовка
        for i, row in enumerate(table.rows[header_row_index + 1:], start=1):
            try:
                # Получаем текст из всех ячеек
                row_texts = [cell.text.strip() for cell in row.cells]
                
                # Фильтруем пустые значения
                row_texts = [text for text in row_texts if text]
                
                if len(row_texts) < 2:
                    continue
                
                # Определяем какие данные где
                # Обычно: [номер, ФИО, код] или [ФИО, код]
                if len(row_texts) >= 3:
                    number = row_texts[0]
                    full_name = row_texts[1]
                    code = row_texts[2]
                elif len(row_texts) == 2:
                    number = str(i)
                    full_name = row_texts[0]
                    code = row_texts[1]
                else:
                    continue
                
                # Проверяем что ФИО похоже на ФИО (содержит буквы)
                if not any(c.isalpha() for c in full_name):
                    continue
                
                # Проверяем что код похож на код (содержит sbph или цифры/буквы и слэши)
                if code and (code.startswith('sbph') or '/' in code or any(c.isalnum() for c in code)):
                    grade8_codes.append({
                        "number": number,
                        "full_name": full_name,
                        "code": code
                    })
                    print(f"DEBUG: Добавлен ученик {full_name}: {code}")
                else:
                    print(f"DEBUG: Пропущена строка - код не похож на код: '{code}'")
                    
            except Exception as e:
                print(f"DEBUG: Ошибка обработки строки {i}: {e}")
                continue
        
        print(f"DEBUG: Всего найдено кодов 8 класса: {len(grade8_codes)}")
        return grade8_codes
    
    def _extract_grade9_codes(self) -> List[str]:
        """
        Извлекает коды для 9 класса из текста документа
        
        Ожидается, что коды идут списком после заголовка типа "Физика за 9 класс"
        """
        grade9_codes = []
        capture_codes = False
        
        print("DEBUG: Начинаем поиск кодов 9 класса...")
        
        # Сначала проверяем параграфы
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # Начинаем захватывать коды после заголовка "... за 9 класс"
            if "9 класс" in text.lower() or "9класс" in text.lower():
                capture_codes = True
                print(f"DEBUG: Найден заголовок 9 класса: '{text}'")
                continue
            
            # Если мы в режиме захвата
            if capture_codes:
                # Проверяем, не начался ли новый раздел
                if text.startswith("---") or "8 класс" in text.lower():
                    print(f"DEBUG: Остановка захвата на строке: '{text}'")
                    break
                
                # Если строка похожа на код (начинается с sbph59 или содержит паттерн кода)
                if text.startswith("sbph59") or (text.startswith("sbph") and "/9/" in text):
                    grade9_codes.append(text)
                    print(f"DEBUG: Найден код 9 класса: {text}")
                elif "/" in text and any(c.isalnum() for c in text) and len(text) > 10:
                    # Возможно это код в другом формате
                    parts = text.split()
                    for part in parts:
                        if "/" in part and len(part) > 10:
                            grade9_codes.append(part)
                            print(f"DEBUG: Найден код 9 класса (альт): {part}")
        
        # Если не нашли в параграфах, проверяем таблицы
        if not grade9_codes:
            print("DEBUG: Коды в параграфах не найдены, проверяем таблицы...")
            
            for table_idx, table in enumerate(self.document.tables):
                # Пропускаем первую таблицу (она с учениками)
                if table_idx == 0:
                    continue
                
                print(f"DEBUG: Проверяем таблицу {table_idx}...")
                
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text.startswith("sbph59") or (text.startswith("sbph") and "/9/" in text):
                            grade9_codes.append(text)
                            print(f"DEBUG: Найден код в таблице: {text}")
        
        print(f"DEBUG: Всего найдено кодов 9 класса: {len(grade9_codes)}")
        return grade9_codes


def parse_olympiad_file(file_path: str) -> Dict:
    """
    Удобная функция для парсинга файла олимпиады
    
    Args:
        file_path: путь к .docx файлу
        
    Returns:
        Dict с распарсенными данными
    """
    parser = OlympiadDocParser(file_path)
    return parser.parse()


# Пример использования
if __name__ == "__main__":
    # Тестовый парсинг
    import sys
    
    if len(sys.argv) < 2:
        print("Использование: python docx_parser.py <файл.docx>")
        sys.exit(1)
    
    result = parse_olympiad_file(sys.argv[1])
    
    print(f"Предмет: {result['subject']}")
    print(f"\n8 класс ({len(result['grade8_codes'])} учеников):")
    for student in result['grade8_codes'][:3]:
        print(f"  {student['number']}. {student['full_name']} - {student['code']}")
    if len(result['grade8_codes']) > 3:
        print("  ...")
    
    print(f"\n9 класс ({len(result['grade9_codes'])} кодов):")
    for code in result['grade9_codes'][:3]:
        print(f"  {code}")
    if len(result['grade9_codes']) > 3:
        print("  ...")
